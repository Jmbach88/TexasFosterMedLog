"""
Export Module
Handles generation of Word documents and PDF conversion
"""

import os
from docx import Document
from docx.shared import Pt
from collections import defaultdict
from typing import Dict, Optional
from .resource_manager import ResourceManager


class ExportManager:
    """Manages export of logs to Word and PDF formats"""

    def __init__(self, template_path: str = None):
        """
        Initialize ExportManager

        Args:
            template_path: Optional path to template (for backward compatibility)
                          If None, uses ResourceManager to find template
        """
        if template_path is None:
            # Use ResourceManager to find bundled template
            resource_manager = ResourceManager()
            self.template_path = str(resource_manager.get_template_path())
        else:
            # Legacy mode for development/testing
            self.template_path = template_path

        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template not found: {self.template_path}")

    def _validate_template(self):
        """Validate template exists and has required structure"""
        # Check template file exists
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template file not found: {self.template_path}")

        # Validate template can be opened and has required tables
        try:
            doc = Document(self.template_path)

            # Check that template has at least 3 tables (for standard medication log)
            if len(doc.tables) < 3:
                raise ValueError(
                    f"Invalid template: Expected at least 3 tables, found {len(doc.tables)}. "
                    "Template may be corrupted or incorrect."
                )

        except Exception as e:
            if isinstance(e, (FileNotFoundError, ValueError)):
                raise
            raise RuntimeError(f"Failed to validate template: {str(e)}")

    def export_to_word(self, profile_data: Dict, log_data: Dict,
                       output_path: str, export_method: str = 'extended_table',
                       include_images: bool = False, med_card_images: list = None) -> str:
        """
        Export medication log to Word document

        Args:
            profile_data: Patient profile information
            log_data: Log data with administration entries
            output_path: Where to save the Word document
            export_method: 'extended_table' or 'continuation_pages'

        Returns:
            Path to generated document
        """
        # Validate template before export
        self._validate_template()

        # Load template
        doc = Document(self.template_path)

        # Fill basic information
        self._fill_basic_info(doc, profile_data, log_data)

        # Fill administration log table based on method
        if export_method == 'continuation_pages':
            self._fill_administration_log_continuation(doc, profile_data, log_data)
        else:
            self._fill_administration_log_extended(doc, log_data['administration_log'])

        # Add medication card images if requested
        if include_images and med_card_images:
            self._add_medication_images(doc, med_card_images)

        # Save document
        doc.save(output_path)

        return output_path
    
    def _set_cell_font_size(self, cell, size=8):
        """Set font size for all text in a cell"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(size)

    def _fill_basic_info(self, doc: Document, profile_data: Dict, log_data: Dict):
        """Fill basic patient and medication information"""

        replacements = {
            "{{ChildName}}": profile_data.get('child_name', ''),
            "{{FosterHome}}": profile_data.get('foster_home', ''),
            "{{DateMY}}": log_data.get('month_year', ''),
            "{{AllergyContras}}": profile_data.get('allergies', ''),
            "{{Prescriber}}": profile_data.get('prescriber_name', ''),
            "{{PrescriberPhone}}": profile_data.get('prescriber_phone', ''),
            "{{Pharmacy}}": profile_data.get('pharmacy', ''),
            "{{PharmacyPhone}}": profile_data.get('pharmacy_phone', ''),
            "{{MedicineName}}": log_data.get('medicine_name', ''),
            "{{Strength}}": log_data.get('strength', ''),
            "{{Dosage}}": log_data.get('dosage', ''),
            "{{ReasonPrescribed}}": log_data.get('reason_prescribed', ''),
            "{{ReasonPRN}}": log_data.get('reason_prn', '')
        }

        # Replace in all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in replacements.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
    
    def _fill_administration_log_extended(self, doc: Document, admin_log: list):
        """Fill the daily administration log table with extended rows for more than 3 administrations"""

        admin_table = doc.tables[2]

        # Group entries by day
        entries_by_day = defaultdict(list)
        for entry in admin_log:
            entries_by_day[entry['day']].append(entry)

        # Find maximum number of administrations per day
        max_admins = max([len(entries) for entries in entries_by_day.values()]) if entries_by_day else 3

        # If we need more than 3 administrations, add rows to the table
        if max_admins > 3:
            self._extend_table_rows(admin_table, max_admins)

        # Fill each day's column
        for day in sorted(entries_by_day.keys()):
            if day < 1 or day > 31:
                continue

            col_idx = day
            entries = entries_by_day[day]

            # Fill all administrations for the day
            for admin_num, entry in enumerate(entries):
                # Calculate row indices: each admin uses 3 rows (time, initials, amount)
                # Starting at row 1, with 1 header row between each set
                base_row = 1 + (admin_num * 4)  # 1, 5, 9, 13, 17, etc.
                time_row = base_row
                initials_row = base_row + 1
                amount_row = base_row + 2

                time_cell = admin_table.rows[time_row].cells[col_idx]
                initials_cell = admin_table.rows[initials_row].cells[col_idx]
                amount_cell = admin_table.rows[amount_row].cells[col_idx]

                time_cell.text = entry.get('time', '')
                initials_cell.text = entry.get('initials', '')
                amount_cell.text = entry.get('amount_remaining', '')

                # Set font size to 8
                self._set_cell_font_size(time_cell)
                self._set_cell_font_size(initials_cell)
                self._set_cell_font_size(amount_cell)

    def _extend_table_rows(self, table, max_admins: int):
        """Add additional row sets to the administration table for more than 3 administrations"""

        # Current structure has 3 admin sets (rows 1-3, 5-7, 9-11)
        # We need to add sets for admin 4, 5, 6, etc.

        for admin_num in range(3, max_admins):
            # Add a header row (like "Admin 4:", "Admin 5:", etc.)
            header_row = table.add_row()
            header_row.cells[0].text = f"Admin {admin_num + 1}:"

            # Add 3 data rows (time, initials, amount)
            time_row = table.add_row()
            time_row.cells[0].text = "Time:"

            initials_row = table.add_row()
            initials_row.cells[0].text = "Initials:"

            amount_row = table.add_row()
            amount_row.cells[0].text = "Amount Remaining:"

    def _fill_administration_log_continuation(self, doc: Document, profile_data: Dict, log_data: Dict):
        """Fill administration log using continuation pages when more than 3 administrations per day"""

        admin_log = log_data['administration_log']
        admin_table = doc.tables[2]

        # Group entries by day
        entries_by_day = defaultdict(list)
        for entry in admin_log:
            entries_by_day[entry['day']].append(entry)

        # Fill first page (first 3 administrations for each day)
        for day in sorted(entries_by_day.keys()):
            if day < 1 or day > 31:
                continue

            col_idx = day
            entries = entries_by_day[day]

            # Fill first 3 administrations
            for admin_num, entry in enumerate(entries[:3]):
                if admin_num == 0:
                    time_row, initials_row, amount_row = 1, 2, 3
                elif admin_num == 1:
                    time_row, initials_row, amount_row = 5, 6, 7
                else:  # admin_num == 2
                    time_row, initials_row, amount_row = 9, 10, 11

                time_cell = admin_table.rows[time_row].cells[col_idx]
                initials_cell = admin_table.rows[initials_row].cells[col_idx]
                amount_cell = admin_table.rows[amount_row].cells[col_idx]

                time_cell.text = entry.get('time', '')
                initials_cell.text = entry.get('initials', '')
                amount_cell.text = entry.get('amount_remaining', '')

                # Set font size to 8
                self._set_cell_font_size(time_cell)
                self._set_cell_font_size(initials_cell)
                self._set_cell_font_size(amount_cell)

        # Check if any day has more than 3 administrations
        needs_continuation = any(len(entries) > 3 for entries in entries_by_day.values())

        if needs_continuation:
            # Create continuation page(s) for administrations 4-6, 7-9, etc.
            page_num = 2
            start_admin = 3

            while needs_continuation:
                # Check if there are any more administrations to process
                has_more = any(len(entries) > start_admin for entries in entries_by_day.values())

                if not has_more:
                    break

                # Add page break
                doc.add_page_break()

                # Add continuation page header
                doc.add_heading(f'Medication Log - Continuation (Page {page_num})', level=2)

                # Load template again for the table structure
                template_doc = Document(self.template_path)
                continuation_table = template_doc.tables[2]._element
                doc._element.body.append(continuation_table)

                # Get the newly added table
                new_table = doc.tables[-1]

                # Fill basic info in new table
                self._fill_basic_info_in_table(new_table, profile_data, log_data)

                # Fill administrations 4-6 (or 7-9, etc.)
                for day in sorted(entries_by_day.keys()):
                    if day < 1 or day > 31:
                        continue

                    col_idx = day
                    entries = entries_by_day[day]

                    # Get the next 3 administrations
                    for offset, entry in enumerate(entries[start_admin:start_admin + 3]):
                        if offset == 0:
                            time_row, initials_row, amount_row = 1, 2, 3
                        elif offset == 1:
                            time_row, initials_row, amount_row = 5, 6, 7
                        else:  # offset == 2
                            time_row, initials_row, amount_row = 9, 10, 11

                        time_cell = new_table.rows[time_row].cells[col_idx]
                        initials_cell = new_table.rows[initials_row].cells[col_idx]
                        amount_cell = new_table.rows[amount_row].cells[col_idx]

                        time_cell.text = entry.get('time', '')
                        initials_cell.text = entry.get('initials', '')
                        amount_cell.text = entry.get('amount_remaining', '')

                        # Set font size to 8
                        self._set_cell_font_size(time_cell)
                        self._set_cell_font_size(initials_cell)
                        self._set_cell_font_size(amount_cell)

                start_admin += 3
                page_num += 1

                # Check if we need yet another page
                needs_continuation = any(len(entries) > start_admin for entries in entries_by_day.values())

    def _fill_basic_info_in_table(self, table, profile_data: Dict, log_data: Dict):
        """Fill basic info in a specific table (for continuation pages)"""
        replacements = {
            "{{ChildName}}": profile_data.get('child_name', ''),
            "{{DateMY}}": log_data.get('month_year', '') + " (cont.)",
        }

        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    def _add_medication_images(self, doc: Document, image_paths: list):
        """Add medication card images to the document"""
        if not image_paths:
            return

        # Add page break before images section
        doc.add_page_break()

        # Add heading
        heading = doc.add_heading('Medication Images', level=1)

        # Add each image
        from docx.shared import Inches

        for i, image_path in enumerate(image_paths):
            if not os.path.exists(image_path):
                continue

            try:
                # Add some spacing before each image (except the first)
                if i > 0:
                    doc.add_paragraph()

                # Add image filename as caption
                filename = os.path.basename(image_path)
                caption = doc.add_paragraph()
                run = caption.add_run(f"Image {i+1}: {filename}")
                run.font.size = Pt(10)
                run.italic = True

                # Add the image (max width 6 inches to fit on page)
                doc.add_picture(image_path, width=Inches(6))

            except Exception as e:
                # If image can't be added, add a note
                error_para = doc.add_paragraph()
                run = error_para.add_run(f"Could not add image: {filename} ({e})")
                run.font.size = Pt(10)
                run.italic = True

    def export_to_pdf(self, word_path: str, pdf_path: Optional[str] = None) -> Optional[str]:
        """
        Convert Word document to PDF
        
        Args:
            word_path: Path to Word document
            pdf_path: Output PDF path (optional, defaults to same name as .docx)
            
        Returns:
            Path to PDF if successful, None if conversion failed
        """
        if pdf_path is None:
            pdf_path = word_path.replace('.docx', '.pdf')
        
        try:
            from docx2pdf import convert
            convert(word_path, pdf_path)
            return pdf_path
        except ImportError:
            print("Warning: docx2pdf not available. PDF conversion skipped.")
            return None
        except Exception as e:
            print(f"Warning: PDF conversion failed: {e}")
            return None
    
    def get_patient_export_dir(self, profile_id: str, data_dir: str = "data/patients") -> str:
        """Get patient-specific export directory"""
        export_dir = os.path.join(data_dir, profile_id, "exports")
        os.makedirs(export_dir, exist_ok=True)
        return export_dir

    def export_log(self, profile_data: Dict, log_data: Dict,
                   output_dir: Optional[str] = None, filename_base: Optional[str] = None,
                   create_pdf: bool = True, export_method: str = 'extended_table',
                   include_images: bool = False,
                   data_dir: str = "data/patients") -> Dict[str, str]:
        """
        Export log to both Word and optionally PDF

        Args:
            profile_data: Patient profile
            log_data: Log data
            output_dir: Directory to save files (if None, uses patient export folder)
            filename_base: Base filename (without extension)
            create_pdf: Whether to also create PDF
            export_method: 'extended_table' or 'continuation_pages'
            data_dir: Base data directory

        Returns:
            Dictionary with 'docx' and optionally 'pdf' paths
        """
        # Use patient export folder if no output_dir specified
        if output_dir is None:
            profile_id = profile_data.get('child_name', '').lower().replace(' ', '_')
            output_dir = self.get_patient_export_dir(profile_id, data_dir)

        os.makedirs(output_dir, exist_ok=True)

        # Generate filename if not provided
        if filename_base is None:
            medicine_name = log_data.get('medicine_name', 'unknown').replace(' ', '_')
            month_year = log_data['month_year'].replace(' ', '_')
            filename_base = f"{medicine_name}_{month_year}"

        # Get medication card images if requested
        med_card_images = []
        if include_images:
            from core.medication_cards import MedicationCardManager
            med_card_manager = MedicationCardManager(data_dir)
            profile_id = profile_data.get('child_name', '').lower().replace(' ', '_')
            medicine_name = log_data.get('medicine_name', '')

            # Get the medication card to find its images
            card = med_card_manager.get_card(profile_id, medicine_name)
            if card:
                # Get all image paths for this medication
                images = card.get('images', [])
                for image_filename in images:
                    image_path = med_card_manager.get_image_path(profile_id, medicine_name, image_filename)
                    if os.path.exists(image_path):
                        med_card_images.append(image_path)

        # Export Word
        word_path = os.path.join(output_dir, f"{filename_base}.docx")
        self.export_to_word(profile_data, log_data, word_path, export_method,
                           include_images, med_card_images)

        result = {'docx': word_path}

        # Export PDF if requested
        if create_pdf:
            pdf_path = self.export_to_pdf(word_path)
            if pdf_path:
                result['pdf'] = pdf_path

        return result
